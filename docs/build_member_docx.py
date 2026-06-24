# -*- coding: utf-8 -*-
"""生成《成员工作细化说明》Word（要点式段落，无圆点列表，中英文无空格）。"""
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUT = Path(__file__).resolve().parent / "成员工作细化说明.docx"


def tight_zh_en(text: str) -> str:
    """去掉中文与英文/数字/路径符号之间的空格。"""
    prev = None
    while prev != text:
        prev = text
        text = re.sub(r"([\u4e00-\u9fff])\s+([A-Za-z0-9_./\\-])", r"\1\2", text)
        text = re.sub(r"([A-Za-z0-9_./\\-])\s+([\u4e00-\u9fff])", r"\1\2", text)
    return text


def set_doc_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)


def run_font(run, bold=False, size=12):
    run.font.name = "宋体"
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_para(doc, text, bold=False, indent_cm=0, align=None):
    text = tight_zh_en(text)
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run(text)
    run_font(run, bold=bold)
    return p


def add_section_heading(doc, text):
    add_para(doc, text, bold=True)


def add_item_block(doc, title: str, body: str):
    """（n）标题：正文… 合并为一段要点。"""
    add_para(doc, f"{title}：{body}")


def add_item_title_then_lines(doc, title, lines):
    """（n）标题 单独一行，下列要点（无圆点）。"""
    add_para(doc, title)
    for line in lines:
        add_para(doc, line, indent_cm=0.74)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = tight_zh_en(h)
        for p in cell.paragraphs:
            for r in p.runs:
                run_font(r, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = tight_zh_en(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    run_font(r)
    doc.add_paragraph()


def main():
    doc = Document()
    set_doc_font(doc)
    for sec in doc.sections:
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(3.17)
        sec.right_margin = Cm(3.17)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("成员工作细化说明")
    run_font(r, bold=True, size=18)

    add_section_heading(doc, "一、项目组成员")
    add_table(
        doc,
        ["姓名", "学号", "角色"],
        [
            ["唐雨薇", "2023132048", "组长、负责后端、训练数据与算法"],
            ["邓晓虹", "2023132042", "前端、演示数据与结果展示、系统测试"],
        ],
    )

    add_section_heading(doc, "二、成员分工")
    add_table(
        doc,
        ["成员", "主要职责"],
        [
            [
                "唐雨薇",
                "项目总体协调与文档；Flask后端与REST API；SQLite业务数据库；"
                "YOLOv5/YOLOv8模型；训练数据集datasets/wangy；DeepSeek LLM；"
                "前后端联调；GitHub：backend/（除演示图外）、datasets/",
            ],
            [
                "邓晓虹",
                "React前端全部页面；UI与交互；投放演练演示图backend/eco_static/garbage/"
                "的整理与替换；检测结果的表格/映射展示（ecoScenarioMap.js）；"
                "历史记录页数据展示；Axios联调；系统测试与演示；"
                "GitHub：frontend/、eco_static/garbage/演示图",
            ],
        ],
    )

    add_section_heading(doc, "三、具体工作内容介绍（成员详细罗列自己工作内容）")

    add_section_heading(doc, "（一）唐雨薇：后端、训练数据与算法")

    add_para(
        doc,
        "1.总结性描述：本人作为组长，负责EcoVision AI平台的后端整体架构、"
        "训练数据与检测算法。将EcoVision视频检测（YOLOv8）与YOLOv5四类垃圾分类"
        "合并为单一Flask应用（backend/app.py）；完成用户登录注册、图片检测、"
        "历史记录、管理后台、DeepSeek轻量LLM等REST服务；维护SQLite业务数据库"
        "（用户、检测历史）及训练数据集datasets/wangy（YOLO格式标注）；"
        "负责模型权重（best.pt、yolov8n.pt、garbage.yaml）及推理联调；"
        "配置CORS、会话、文件上传等。演示用轮播图片由邓晓虹负责整理，"
        "本人负责静态目录挂载与接口提供。同时负责项目文档、目录规范及"
        "GitHub后端与训练数据上传。",
    )

    add_para(doc, "2.详细工作内容", bold=True)

    add_item_block(
        doc,
        "（1）后端入口与路由",
        "编写/维护backend/app.py：Flask应用初始化、CORS、静态资源（eco_static）、"
        "SPA托管、EcoVision视频流/video_feed、Legacy页面等。"
        "挂载子模块：garbage_api（垃圾分类API）、llm_api（DeepSeek对话与检测解读）。",
    )
    add_item_block(
        doc,
        "（2）垃圾分类API与数据库（backend/garbage_api/extension.py）",
        "设计SQLAlchemy模型：User、Image、DetectionHistory、SystemLog。"
        "实现接口：POST /api/login、/api/register、/api/logout；POST /api/detect；"
        "GET /api/user/history；GET /api/admin/*。"
        "配置上传目录uploads/garbage/，处理multipart上传与中文路径兼容。",
    )
    add_item_block(
        doc,
        "（3）DeepSeek LLM模块（backend/llm_api/）",
        "deepseek_client.py：封装DeepSeek Chat Completions API。"
        "extension.py：提供GET /api/llm/status、POST /api/llm/chat、"
        "POST /api/llm/explain-detection。"
        "通过根目录.env读取DEEPSEEK_API_KEY，密钥不提交Git。",
    )
    add_item_title_then_lines(
        doc,
        "（4）训练数据与业务数据库",
        [
            "SQLite：backend/instance/garbage_classification.db（用户、上传记录、检测历史）。",
            "训练集：datasets/wangy/（约300张图+YOLO标注）。",
            "模型配置：backend/garbage_yolov5/garbage.yaml。",
            "维护requirements.txt、run_backend.bat、.env.example、.gitignore。",
        ],
    )
    add_item_title_then_lines(
        doc,
        "（5）检测算法与模型",
        [
            "backend/waste_classifier.py：YOLOv8实时视频检测。",
            "backend/garbage_yolov5/：best.pt、garbage.yaml、训练/检测脚本。",
            "garbage_api/extension.py中detect_garbage_image()：图片四类推理与画框。",
            "datasets/wangy/：训练数据集（images+labels，YOLO格式）。",
            "backend/training/、train_waste_model.py：YOLOv8侧训练脚本（可选扩展）。",
        ],
    )
    add_item_title_then_lines(
        doc,
        "（6）联调与工程化",
        [
            "与邓晓虹约定接口格式（JSON、result_url、Cookie鉴权）。",
            "联调/api/detect返回的class_name、置信度、bbox与前端展示、投放演练映射。",
            "将backend/、datasets/推送至GitHub。",
        ],
    )

    add_section_heading(doc, "（二）邓晓虹·前端、演示数据与结果展示")

    add_para(
        doc,
        "1.总结性描述：本人负责EcoVision AI平台的React前端、演示数据与检测结果的展示，"
        "以及系统测试。除完成全部页面与交互外，还负责投放演练用演示图片的整理与维护"
        "（backend/eco_static/garbage/四类各1～3张）；负责将接口返回的类别、置信度、"
        "边框在识别工作台以表格展示，并通过ecoScenarioMap.js将YOLOv5结果映射到四色垃圾桶；"
        "负责历史记录页检测数据的列表展示。通过Axios与唐雨薇的后端API联调，"
        "统一UI样式，并完成全流程测试与课堂演示。",
    )

    add_para(doc, "2.详细工作内容", bold=True)

    add_item_title_then_lines(
        doc,
        "（1）工程结构与路由（frontend/src/App.js）",
        [
            "配置React Router：未登录跳转登录页，登录后展示主导航与各功能页。",
            "路由包括：/、/detect、/video-live、/scenario-carousel、/assistant、"
            "/history、/profile、/admin、/about等。",
        ],
    )
    add_item_title_then_lines(
        doc,
        "（2）公共组件",
        [
            "components/Navbar.js：顶部导航。",
            "components/Footer.js、ProtectedRoute.js：页脚与权限路由。",
        ],
    )
    add_item_title_then_lines(
        doc,
        "（3）接口层（frontend/src/services/）",
        [
            "api.js：Axios实例，withCredentials: true。",
            "authService.js：登录、登出、获取当前用户。",
        ],
    )

    add_para(doc, "（4）核心页面实现")
    add_table(
        doc,
        ["页面文件", "功能"],
        [
            ["LoginPage.js/RegisterPage.js", "登录注册"],
            ["HomePage.js", "首页双入口"],
            ["DetectPage.js", "上传图片→/api/detect→展示原图/标注图/表格；AI解读"],
            ["VideoEcoPage.js", "订阅/video_feed实时视频"],
            ["ScenarioCarouselPage.js", "轮播→自动检测→四桶选择→对错判定+语音"],
            ["AssistantPage.js", "Eco助手，/api/llm/chat"],
            ["HistoryPage.js等", "历史、个人信息、管理后台"],
        ],
    )

    add_item_title_then_lines(
        doc,
        "（5）演示数据与业务展示数据",
        [
            "维护backend/eco_static/garbage/四类演示图（recyclable、kitchen、harmful、other）。",
            "DetectPage.js：将/api/detect返回结果渲染为表格。",
            "ecoScenarioMap.js：模型输出映射为投放演练四色桶ID。",
            "HistoryPage.js：展示用户往期检测记录。",
            "resolveApiUrl.js：拼接后端返回的结果图URL。",
        ],
    )
    add_item_title_then_lines(
        doc,
        "（6）工具与交互",
        [
            "scenarioSpeech.js：浏览器中文语音播报。",
            "index.css：全局样式、投放演练、聊天面板等。",
        ],
    )
    add_item_title_then_lines(
        doc,
        "（7）系统测试与协作",
        [
            "功能测试：登录、上传识别、投放演练、Eco助手、历史记录、实时视频等全流程。",
            "与唐雨薇联调FormData上传、401提示、代理排查等。",
            "记录测试用例与问题反馈，配合答辩演示脚本。",
            "将frontend/源码及backend/eco_static/garbage/演示图上传GitHub。",
        ],
    )

    doc.save(OUT)
    print(f"已生成: {OUT}")


if __name__ == "__main__":
    main()
